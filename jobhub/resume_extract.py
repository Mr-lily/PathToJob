# -*- coding: utf-8 -*-
"""简历文件文本提取：PDF / Word(.docx) / 纯文本。

serve.py 上传简历文件后调用本模块提取纯文本。
pdf 用 pypdf，docx 用 python-docx（装在隔离 venv）。
"""

import os


def extract_pdf(path):
    from pypdf import PdfReader
    r = PdfReader(path)
    parts = []
    for page in r.pages:
        try:
            parts.append(page.extract_text() or '')
        except Exception:
            pass
    return '\n'.join(parts)


def extract_docx(path):
    import docx
    d = docx.Document(path)
    lines = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tb in d.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(' | '.join(cells))
    return '\n'.join(lines)


def extract_file(path):
    """按扩展名提取文本。支持 .txt/.pdf/.docx/.doc(老格式走docx尝试)。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_pdf(path)
    if ext in ('.docx', '.doc'):
        # .doc 老格式 python-docx 打不开；报错让前端提示转存 .docx
        return extract_docx(path)
    # txt / 其他：按 utf-8 / gbk 尝试读
    for enc in ('utf-8', 'gbk', 'utf-8-sig'):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(path, encoding='utf-8', errors='ignore') as f:
        return f.read()
